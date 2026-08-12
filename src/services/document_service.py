"""Safe, format-specific document extraction and ESG enrichment."""

from __future__ import annotations

import asyncio
import csv
import hashlib
import io
import json
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from uuid import NAMESPACE_URL, uuid5

from langchain_core.documents import Document
from langchain_text_splitters import MarkdownTextSplitter, RecursiveCharacterTextSplitter
from langsmith import traceable

from ..core.config import settings
from ..core.exceptions import DocumentProcessingError
from ..core.logging import LoggingMixin
from ..core.taxonomy import get_taxonomy


class DocumentProcessor(LoggingMixin):
    """Extract text using small, dedicated libraries rather than a shell pipeline."""

    def __init__(self) -> None:
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        self.markdown_splitter = MarkdownTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )

    @traceable(name="process_document")
    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[Document]:
        safe_filename = Path(filename or "").name
        if not safe_filename or safe_filename in {".", ".."}:
            raise DocumentProcessingError("A valid filename is required")
        if not file_content:
            raise DocumentProcessingError(f"{safe_filename} is empty")
        if len(file_content) > settings.max_file_size:
            raise DocumentProcessingError(
                f"{safe_filename} exceeds the {settings.max_document_size_mb:g} MB upload limit"
            )

        extension = Path(safe_filename).suffix.lower()
        if extension not in settings.supported_extensions:
            raise DocumentProcessingError(
                f"Unsupported file extension '{extension or '(none)'}'. "
                f"Allowed: {', '.join(settings.supported_extensions)}"
            )

        try:
            source_documents = await asyncio.to_thread(
                self._extract_documents,
                extension,
                file_content,
                safe_filename,
            )
            splitter = self.markdown_splitter if extension == ".md" else self.text_splitter
            chunks = splitter.split_documents(source_documents)
        except DocumentProcessingError:
            raise
        except Exception as exc:
            self.logger.exception("Document extraction failed", filename=safe_filename)
            raise DocumentProcessingError(f"Could not extract text from {safe_filename}") from exc

        chunks = [chunk for chunk in chunks if chunk.page_content.strip()]
        if not chunks:
            raise DocumentProcessingError(f"No readable text was found in {safe_filename}")
        if len(chunks) > settings.max_chunks_per_document:
            raise DocumentProcessingError(
                f"{safe_filename} produced {len(chunks)} chunks; the configured limit is "
                f"{settings.max_chunks_per_document}"
            )

        document_hash = hashlib.sha256(file_content).hexdigest()
        mime_type = self._mime_type(extension)
        base_metadata = {
            "filename": safe_filename,
            "mime_type": mime_type,
            "file_size_mb": round(len(file_content) / (1024 * 1024), 6),
            "document_hash": document_hash,
            **(metadata or {}),
        }
        for chunk in chunks:
            chunk.metadata.update(base_metadata)
        return chunks

    def _extract_documents(
        self, extension: str, file_content: bytes, filename: str
    ) -> list[Document]:
        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".txt": self._extract_text,
            ".md": self._extract_text,
            ".csv": self._extract_csv,
            ".xlsx": self._extract_xlsx,
        }
        return extractors[extension](file_content, filename)

    @staticmethod
    def _extract_text(file_content: bytes, filename: str) -> list[Document]:
        try:
            text = file_content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentProcessingError(f"{filename} must be UTF-8 encoded") from exc
        return [Document(page_content=text, metadata={"source": filename})]

    @staticmethod
    def _extract_pdf(file_content: bytes, filename: str) -> list[Document]:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_content), strict=False)
        return [
            Document(
                page_content=page.extract_text() or "",
                metadata={"source": filename, "page": page_number},
            )
            for page_number, page in enumerate(reader.pages, start=1)
        ]

    @staticmethod
    def _extract_docx(file_content: bytes, filename: str) -> list[Document]:
        from docx import Document as WordDocument

        word_document = WordDocument(io.BytesIO(file_content))
        paragraphs = [paragraph.text for paragraph in word_document.paragraphs if paragraph.text]
        for table in word_document.tables:
            paragraphs.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return [Document(page_content="\n".join(paragraphs), metadata={"source": filename})]

    @staticmethod
    def _extract_csv(file_content: bytes, filename: str) -> list[Document]:
        try:
            text = file_content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentProcessingError(f"{filename} must be UTF-8 encoded") from exc
        reader = csv.DictReader(io.StringIO(text))
        rows = [json.dumps(row, ensure_ascii=False) for row in reader]
        return [Document(page_content="\n".join(rows), metadata={"source": filename})]

    @staticmethod
    def _extract_xlsx(file_content: bytes, filename: str) -> list[Document]:
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
        documents: list[Document] = []
        try:
            for worksheet in workbook.worksheets:
                rows = [
                    " | ".join("" if value is None else str(value) for value in row)
                    for row in worksheet.iter_rows(values_only=True)
                ]
                documents.append(
                    Document(
                        page_content="\n".join(rows),
                        metadata={"source": filename, "sheet": worksheet.title},
                    )
                )
        finally:
            workbook.close()
        return documents

    @staticmethod
    def _mime_type(extension: str) -> str:
        return {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".csv": "text/csv",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        }[extension]


class DocumentService(LoggingMixin):
    def __init__(self, processor: DocumentProcessor | None = None) -> None:
        self.processor = processor or DocumentProcessor()

    @traceable(name="process_uploaded_document")
    async def process_uploaded_document(
        self,
        file_content: bytes,
        filename: str,
        esg_framework: str | None = None,
        document_type: str | None = None,
        company_id: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> list[Document]:
        taxonomy = get_taxonomy()
        if esg_framework and esg_framework not in taxonomy.framework_ids:
            raise DocumentProcessingError(f"Unknown ESG framework: {esg_framework}")
        if document_type and document_type not in taxonomy.document_types:
            raise DocumentProcessingError(f"Unknown document type: {document_type}")

        enriched_metadata: dict[str, Any] = {
            "processed_at": datetime.now(UTC).isoformat(),
            **(metadata or {}),
        }
        for key, value in {
            "esg_framework": esg_framework,
            "document_type": document_type,
            "company_id": company_id,
        }.items():
            if value:
                enriched_metadata[key] = value

        documents = await self.processor.process_document(
            file_content,
            filename,
            enriched_metadata,
        )
        document_hash = str(documents[0].metadata["document_hash"])
        for index, document in enumerate(documents):
            document.metadata["chunk_index"] = index
            document.metadata["chunk_id"] = str(
                uuid5(NAMESPACE_URL, f"esg:{document_hash}:{index}")
            )
            if not document.metadata.get("esg_category"):
                category = self._detect_esg_category(document.page_content)
                if category:
                    document.metadata["esg_category"] = category
            document.metadata = self._clean_metadata(document.metadata)
        return documents

    def _detect_esg_category(self, content: str) -> str | None:
        content_lower = content.casefold()
        scores = {
            category: sum(keyword.casefold() in content_lower for keyword in keywords)
            for category, keywords in get_taxonomy().categories.items()
        }
        best_category, best_score = max(scores.items(), key=lambda item: item[1])
        return best_category if best_score else None

    async def batch_process_documents(
        self,
        files: list[tuple[bytes, str]],
        **common_metadata: Any,
    ) -> tuple[list[Document], list[str]]:
        results = await asyncio.gather(
            *(
                self.process_uploaded_document(content, filename, **common_metadata)
                for content, filename in files
            ),
            return_exceptions=True,
        )
        documents: list[Document] = []
        errors: list[str] = []
        for (_, filename), result in zip(files, results, strict=True):
            if isinstance(result, BaseException):
                errors.append(f"{filename}: {result}")
            else:
                documents.extend(result)
        return documents, errors

    @staticmethod
    def _clean_metadata(metadata: dict[str, Any]) -> dict[str, str | int | float | bool]:
        return {
            key: value if isinstance(value, (str, int, float, bool)) else str(value)
            for key, value in metadata.items()
            if value is not None
        }


document_service = DocumentService()
